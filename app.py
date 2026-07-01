# C:\aka\app.py
# AKA v3.9.9 - Final Production with PERFECT KOP SURAT
# QA: Gemini + DeepSeek | UI: Meta | Final: LUCA

import streamlit as st
import pandas as pd
import numpy as np
import io
import datetime
import openpyxl
import re
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.comments import Comment
from openpyxl.utils import get_column_letter
from contextlib import closing
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- TERBILANG DENGAN FALLBACK ---
try:
    from num2words import num2words
    def terbilang_rupiah(nilai):
        try:
            # Batasi hanya integer untuk menghindari error desimal
            words = num2words(int(nilai), lang='id').capitalize()
            return f"{words} Rupiah"
        except:
            return f"Rp {nilai:,.0f} (terbilang otomatis tidak tersedia)"
except ImportError:
    def terbilang_rupiah(nilai):
        return f"Rp {nilai:,.0f}"

st.set_page_config(page_title="AKA Validator Pro 2026", page_icon="⚖️", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.block-container { padding-top: 2rem; padding-bottom: 0rem; }
.main-header { font-size: 42px!important; font-weight: 700; text-align: center; padding: 10px 0px; background: linear-gradient(135deg, #8B5CF6, #6366F1, #06B6D4); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.sub-header { text-align: center; opacity: 0.7; font-size: 16px; margin-bottom: 30px; }
.upload-box { border: 3px dashed #8B5CF6; border-radius: 20px; padding: 30px; text-align: center; max-width: 700px; margin: 0 auto 20px auto; background: rgba(139, 92, 246, 0.05); }
.card-modern { padding: 20px 15px; border-radius: 16px; text-align: center; border: 1px solid rgba(128, 128, 128, 0.1); background: rgba(255, 255, 255, 0.05); box-shadow: 0 4px 12px rgba(0,0,0,0.1); }
.card-modern h4 { margin: 0; font-weight: 500; font-size: 0.9rem; opacity: 0.7; }
.card-modern h2 { margin: 5px 0 0; font-weight: 700; font-size: 1.8rem; }
.bg-green { background: linear-gradient(135deg, #10B981 0%, #059669 100%); color: white; }
.bg-red { background: linear-gradient(135deg, #EF4444 0%, #DC2626 100%); color: white; }
.bg-orange { background: linear-gradient(135deg, #F59E0B 0%, #D97706 100%); color: white; }
.bg-blue { background: linear-gradient(135deg, #3B82F6 0%, #2563EB 100%); color: white; }
.stButton>button,.stDownloadButton>button { background: linear-gradient(45deg, #8B5CF6 0%, #6366F1 100%); color: white; border-radius: 12px; padding: 12px 28px; border: none; font-weight: 600; width: 100%; }
.legal-note { background: rgba(245,158,11,0.1); border-left: 4px solid #F59E0B; padding: 12px; border-radius: 8px; margin-top: 10px; font-size: 0.9rem; }
</style>
""", unsafe_allow_html=True)

st.markdown('<p class="main-header">⚖️ AKA Validator Pro 2026</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Agentic Audit + Compliance Check | Perpres 12/2021 & SE PUPR 07/2023 🔍</p>', unsafe_allow_html=True)

st.markdown('<div class="upload-box">', unsafe_allow_html=True)
uploaded_file = st.file_uploader("📂 Upload File RAB Penawar (.xlsx)", type=['xlsx', 'xls'], label_visibility="collapsed")
st.markdown('</div>', unsafe_allow_html=True)

# --- FUNGSI CORE ---
def unmerge_and_fill(worksheet):
    for merged_cell in list(worksheet.merged_cells.ranges):
        min_col, min_row, max_col, max_row = merged_cell.bounds
        top_left_value = worksheet.cell(row=min_row, column=min_col).value
        try: worksheet.unmerge_cells(str(merged_cell))
        except: continue
        for row in worksheet.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            for cell in row: cell.value = top_left_value
    return worksheet

def score_sheet_for_rab(ws):
    score = 0
    text_blob = ' '.join([str(cell.value).lower() for row in ws.iter_rows(max_row=15) for cell in row if cell.value])
    rab_keywords = ['uraian', 'volume', 'harga satuan', 'total harga', 'jumlah', 'satuan', 'vol', 'qty']
    score += sum(10 for kw in rab_keywords if kw in text_blob)
    return score

def detect_columns_and_header(ws):
    rows = list(ws.iter_rows(values_only=True))
    col_map = {}
    keyword_map = {
        'no_col': ['no', 'nomor', '#', 'item'],
        'vol_col': ['volume', 'vol', 'qty', 'kuantitas'],
        'unit_col': ['satuan', 'sat', 'unit'],
        'total_col': ['total harga', 'jumlah harga', 'subtotal', 'total', 'jml harga'],
        'price_col': ['harga satuan', 'harsat', 'unit price', 'harga'],
        'desc_col': ['uraian', 'deskripsi', 'pekerjaan', 'nama item']
    }
    best_match_count = 0
    header_row_idx = None
    for i, row in enumerate(rows[:20]):
        matches = 0; temp_map = {}; used_columns = set()
        for key, keywords in keyword_map.items():
            for col_idx, cell_value in enumerate(row):
                if col_idx in used_columns: continue
                if cell_value:
                    cell_str = str(cell_value).strip().lower()
                    if any(kw == cell_str or (len(kw) > 4 and kw in cell_str) for kw in keywords):
                        temp_map[key] = col_idx; used_columns.add(col_idx); matches += 1; break
        if matches >= 3 and matches > best_match_count:
            best_match_count = matches; header_row_idx = i; col_map = temp_map
    return col_map, header_row_idx

def clean_indonesian_number(val):
    if val is None: return 0
    if isinstance(val, (int, float)): return val
    val_str = str(val).replace('Rp', '').replace('rp', '').replace(' ', '').strip()
    val_str = re.sub(r'[^\d,.]', '', val_str)
    val_str = re.sub(r'[.,]$', '', val_str)
    val_str = val_str.replace('.', '').replace(',', '.')
    return pd.to_numeric(val_str, errors='coerce')

def analyze_root_cause(row):
    notes = []
    if row['__is_error']:
        if row['__vol'] == 0 and row['__total_original'] > 0: notes.append("ITEM SILUMAN: Vol 0 tapi ada harga.")
        elif row['__vol'] > 0 and row['__price'] == 0 and row['__total_original'] > 0: notes.append("HARGA NOL: Volume ada tapi harga satuan 0.")
        else: notes.append("SALAH KETIK/RUMUS: Hasil kali tidak sama.")
    if row['__is_subtotal']: notes.append("BARIS SUBTOTAL: Tidak diaudit aritmatik.")
    return " | ".join(notes) if notes else "Aman"

def process_agentic_audit(file_bytes, selected_sheet):
    with closing(openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)) as wb:
        ws = wb[selected_sheet]; ws = unmerge_and_fill(ws)
        col_map, header_row_idx = detect_columns_and_header(ws)
        if header_row_idx is None or col_map.get('vol_col') is None or col_map.get('price_col') is None:
            raise ValueError("Gagal memetakan struktur tabel RAB.")
        rows = list(ws.iter_rows(values_only=True)); records = []
        for idx, row in enumerate(rows[header_row_idx + 1:], start=header_row_idx + 2):
            if not any(v is not None for v in row): continue
            records.append({'excel_row_num': idx,'raw_desc': row[col_map['desc_col']] if col_map.get('desc_col') is not None and col_map['desc_col'] < len(row) else '','raw_vol': row[col_map['vol_col']] if col_map['vol_col'] < len(row) else 0,'raw_price': row[col_map['price_col']] if col_map['price_col'] < len(row) else 0,'raw_total': row[col_map['total_col']] if col_map.get('total_col') is not None and col_map['total_col'] < len(row) else 0})
        df = pd.DataFrame(records)
        if df.empty: raise ValueError("Tidak ditemukan data item pekerjaan.")
        df['__vol'] = df['raw_vol'].apply(clean_indonesian_number).fillna(0)
        df['__price'] = df['raw_price'].apply(clean_indonesian_number).fillna(0)
        df['__total_original'] = df['raw_total'].apply(clean_indonesian_number).fillna(0)
        df['__is_subtotal'] = df['raw_desc'].astype(str).fillna('').str.lower().str.contains('total|jumlah|sub|rekap') | ((df['__vol'] == 0) & (df['__price'] == 0) & (df['__total_original'] > 0))
        df['__total_correct'] = np.where(df['__is_subtotal'], df['__total_original'], df['__vol'] * df['__price'])
        df['__selisih'] = np.where(df['__is_subtotal'], 0, df['__total_correct'] - df['__total_original'])
        df['__is_error'] = (~df['__is_subtotal']) & ((abs(df['__selisih']) > 0.5) | ((df['__vol'] > 0) & (df['__price'] == 0) & (df['__total_original'] > 0)))
        df['__is_anomaly'] = (~df['__is_subtotal']) & (df['__vol'] == 0) & (df['__total_original'] > 0)
        df['__keterangan'] = df.apply(analyze_root_cause, axis=1)
    return df, col_map

def compliance_check(total_penawaran_koreksi, total_hps, ada_temuan):
    if total_hps == 0:
        return {"status": "TANPA HPS", "dasar_hukum": "Input HPS untuk evaluasi", "tindak_lanjut": "Masukkan nilai HPS.", "persentase": 0, "warna": "bg-blue", "action": None}
    persentase = (total_penawaran_koreksi / total_hps) * 100
    if total_penawaran_koreksi < total_hps * 0.8:
        return {"status": "TIDAK WAJAR", "dasar_hukum": "Perpres 12/2021 Ps. 48: <80% HPS", "tindak_lanjut": "WAJIB Klarifikasi, jika gagal GUGUR.", "persentase": persentase, "warna": "bg-red", "action": "klarifikasi"}
    elif persentase > 110:
        return {"status": "INDIKASI KEMAHALAN", "dasar_hukum": "SE PUPR 07/2023: >110% HPS", "tindak_lanjut": "Wajib negosiasi efisiensi.", "persentase": persentase, "warna": "bg-orange", "action": "negosiasi"}
    elif ada_temuan > 0:
        return {"status": "WAJAR DENGAN CATATAN", "dasar_hukum": "Perpres 12/2021 Ps.60: Koreksi Aritmatik", "tindak_lanjut": f"Terdapat {ada_temuan} temuan. Harga kontrak pakai nilai terkoreksi.", "persentase": persentase, "warna": "bg-orange", "action": None}
    else:
        return {"status": "WAJAR", "dasar_hukum": "Memenuhi ambang batas", "tindak_lanjut": "Lanjut evaluasi teknis.", "persentase": persentase, "warna": "bg-green", "action": None}

def auto_fit_columns(worksheet):
    max_rows_to_scan = 100
    for col in worksheet.columns:
        max_len = 0
        for idx, cell in enumerate(col):
            if idx > max_rows_to_scan: break
            if cell.value:
                for line in str(cell.value).split('\n'): max_len = max(max_len, len(line))
        worksheet.column_dimensions[get_column_letter(col[0].column)].width = max(max_len + 3, 12)

def generate_preserved_excel(file_bytes, selected_sheet, df_result, col_map, compliance, total_hps):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=False)
    ws = wb[selected_sheet]
    red_fill = PatternFill(start_color="FDE8E8", end_color="FDE8E8", fill_type="solid"); red_font = Font(color="9B1C1C", bold=True)
    col_total_idx = col_map.get('total_col')
    if col_total_idx is not None: col_total_idx += 1
    errors_log = []
    if col_total_idx:
        for _, r in df_result.iterrows():
            if r['__is_error'] or r['__is_anomaly']:
                c = ws.cell(row=int(r['excel_row_num']), column=col_total_idx); c.fill = red_fill; c.font = red_font; c.comment = Comment(f"⚠ {r['__keterangan']}", "AKA v3.9.9")
                errors_log.append([r['excel_row_num'], r['raw_desc'], r['__total_original'], r['__total_correct'], r['__selisih'], r['__keterangan']])
    if "Audit_Trail" in wb.sheetnames: del wb["Audit_Trail"]
    ws_a = wb.create_sheet("Audit_Trail"); ws_a.append(["Baris","Uraian","Asli","Koreksi","Selisih","Analisa"])
    for e in errors_log: ws_a.append(e)
    if "Berita_Acara" in wb.sheetnames: del wb["Berita_Acara"]
    ws_b = wb.create_sheet("Berita_Acara", 0)
    df_items = df_result[~df_result['__is_subtotal']]; total_asli = df_items['__total_original'].sum(); total_kor = df_items['__total_correct'].sum()
    for row in [["BERITA ACARA KOREKSI ARITMATIK"],[f"Tanggal: {datetime.date.today()}"],[""],["Hasil:"],[f"Total Asli: Rp {total_asli:,.0f}"],[f"Total Koreksi: Rp {total_kor:,.0f}"],[f"Selisih: Rp {total_kor-total_asli:,.0f}"],[""],[f"Status: {compliance['status']}"],[f"Dasar: {compliance['dasar_hukum']}"],[f"Tindak Lanjut: {compliance['tindak_lanjut']}"]]: ws_b.append(row)
    ws_b['A1'].font = Font(bold=True, size=14); ws_b.merge_cells('A1:F1'); ws_b['A1'].alignment = Alignment(horizontal='center')
    for sn in wb.sheetnames: auto_fit_columns(wb[sn])
    out = io.BytesIO(); wb.save(out); wb.close()
    return out.getvalue()

# =========================================================
# PERBAIKAN UTAMA: FITUR SURAT DINAS DENGAN KOP SURAT 100% RAPI
# =========================================================
def add_kop_pu(doc):
    # PAKAI IMAGE KOP
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Ubah lebar menjadi 6 inches (agar pas di margin A4 standar)
        run.add_picture("kop_pu_header.png", width=Inches(6.0))
        p.paragraph_format.space_after = Pt(2) # Kurangi jarak bawah gambar
        # Tambahkan garis pemisah tipis di bawah kop
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 80) # Garis pemisah
        run.font.size = Pt(2) # Buat garis sangat tipis
        p.paragraph_format.space_after = Pt(6)
    except Exception:
        # Fallback jika gambar tidak ditemukan
        p = doc.add_paragraph("KEMENTERIAN PEKERJAAN UMUM DAN PERUMAHAN RAKYAT")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.bold = True
        p.paragraph_format.space_after = Pt(12)

def surat_klarifikasi_docx(total, hps, persen):
    doc = Document()
    add_kop_pu(doc)
    
    # PERBAIKAN: Gunakan Paragraph, BUKAN add_heading (level 0). Ini yang bikin jarak terlalu jauh!
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT UNDANGAN KLARIFIKASI HARGA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph(f'Nomor: .../KLAR/{datetime.date.today().year}')
    doc.add_paragraph('Kepada Yth.\nPenyedia\nDi Tempat')
    
    p = doc.add_paragraph()
    p.add_run(f'Berdasarkan evaluasi, penawaran Saudara sebesar Rp {total:,.0f} ').bold = True
    p.add_run(f'({terbilang_rupiah(total)}) ').italic = True
    p.add_run(f'atau {persen:.1f}% dari HPS Rp {hps:,.0f} ').bold = True
    p.add_run('di bawah 80% HPS.')
    
    doc.add_paragraph('Sesuai Perpres 12/2021 Pasal 48, Saudara WAJIB hadir klarifikasi dengan membawa:')
    for item in ['Analisa Harga Satuan detail','Bukti harga material (invoice/quotation)','Surat dukungan pabrik/distributor','Pernyataan kesanggupan bermaterai']:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_paragraph('Jika tidak dapat membuktikan kewajaran dalam 3 hari kerja, penawaran dinyatakan GUGUR.').runs[0].bold = True
    doc.add_paragraph('\n\nPalangka Raya, ' + datetime.date.today().strftime('%d %B %Y') + '\n\nPejabat Pengadaan')
    
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def surat_negosiasi_docx(total, hps, persen):
    doc = Document()
    add_kop_pu(doc)
    
    # PERBAIKAN: Gunakan Paragraph
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SURAT UNDANGAN NEGOSIASI HARGA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    
    doc.add_paragraph(f'Nomor: .../NEGO/{datetime.date.today().year}')
    doc.add_paragraph('Kepada Yth.\nPenyedia\nDi Tempat')
    
    p = doc.add_paragraph()
    p.add_run(f'Penawaran Saudara Rp {total:,.0f} ').bold = True
    p.add_run(f'({terbilang_rupiah(total)}) ').italic = True
    p.add_run(f'atau {persen:.1f}% dari HPS Rp {hps:,.0f} ').bold = True
    p.add_run('melebihi 110% HPS.')
    
    doc.add_paragraph('Sesuai SE PUPR 07/2023, Saudara diundang untuk negosiasi efisiensi anggaran.')
    doc.add_paragraph('\n\nPalangka Raya, ' + datetime.date.today().strftime('%d %B %Y') + '\n\nPejabat Pengadaan')
    
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()
# =========================================================

if uploaded_file:
    bytes_data = uploaded_file.getvalue()
    with closing(openpyxl.load_workbook(io.BytesIO(bytes_data), data_only=False)) as wb_temp:
        sheets = wb_temp.sheetnames; scores = {n: score_sheet_for_rab(wb_temp[n]) for n in sheets}; best = max(scores, key=scores.get)
    st.markdown("### ⚙ Konfigurasi Audit")
    c1, c2 = st.columns(2)
    with c1: st.info(f"💡 Rekomendasi: `{best}`", icon="🤖"); sel_sheet = st.selectbox("Pilih Sheet:", sheets, index=sheets.index(best))
    with c2: hps = st.number_input("Input Total HPS (Rp)", min_value=0, value=0, step=1000000)
    if st.button("🚀 Mulai Audit", use_container_width=True):
        with st.spinner('Audit...'):
            try:
                df, cmap = process_agentic_audit(bytes_data, sel_sheet)
                df_items = df[~df['__is_subtotal']]
                if len(df_items) == 0: st.error("❌ Tidak ada item detail ditemukan. Pilih sheet RAB yang benar."); st.stop()
                total_asli = df_items['__total_original'].sum(); total_kor = df_items['__total_correct'].sum(); selisih = total_kor - total_asli; item_err = df_items['__is_error'].sum() + df_items['__is_anomaly'].sum()
                comp = compliance_check(total_kor, hps, item_err)
                st.divider(); st.subheader("⚖ Evaluasi Kepatuhan")
                if comp["status"] == "TANPA HPS": st.markdown(f'<div class="card-modern {comp["warna"]}"><h4>{comp["status"]}</h4><h2>HPS Belum Diinput</h2><small>{comp["dasar_hukum"]}</small></div>', unsafe_allow_html=True)
                else: st.markdown(f'<div class="card-modern {comp["warna"]}"><h4>{comp["status"]}</h4><h2>{comp["persentase"]:.1f}% HPS</h2><small>{comp["dasar_hukum"]}</small></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="legal-note"><b>Tindak Lanjut:</b> {comp["tindak_lanjut"]}</div>', unsafe_allow_html=True)

                if comp["action"] == "klarifikasi":
                    docx_bytes = surat_klarifikasi_docx(total_kor, hps, comp["persentase"])
                    st.download_button("📄 Download Surat Klarifikasi (.docx)", docx_bytes, "Surat_Klarifikasi.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_klar")
                if comp["action"] == "negosiasi":
                    docx_bytes = surat_negosiasi_docx(total_kor, hps, comp["persentase"])
                    st.download_button("📄 Download Surat Negosiasi (.docx)", docx_bytes, "Surat_Negosiasi.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_nego")

                st.divider()
                k1,k2,k3,k4 = st.columns(4)
                k1.markdown(f'<div class="card-modern bg-blue"><h4>Item</h4><h2>{len(df_items)}</h2></div>', unsafe_allow_html=True)
                k2.markdown(f'<div class="card-modern bg-red"><h4>Temuan</h4><h2>{item_err}</h2></div>', unsafe_allow_html=True)
                k3.markdown(f'<div class="card-modern bg-green"><h4>Koreksi</h4><h2>Rp {total_kor:,.0f}</h2></div>', unsafe_allow_html=True)
                k4.markdown(f'<div class="card-modern bg-red"><h4>Selisih</h4><h2>Rp {selisih:,.0f}</h2></div>', unsafe_allow_html=True)

                st.divider()
                err_df = df_items[(df_items['__is_error'])|(df_items['__is_anomaly'])]
                if not err_df.empty:
                    disp = err_df[['excel_row_num','raw_desc','raw_vol','raw_price','raw_total','__total_correct','__selisih','__keterangan']].copy()
                    disp.columns = ['Baris','Uraian','Vol','Harga','Asli','Koreksi','Selisih','Analisa']
                    st.dataframe(disp, use_container_width=True, hide_index=True)
                else: st.success("✅ Aman!")
                out = generate_preserved_excel(bytes_data, sel_sheet, df, cmap, comp, hps)
                st.download_button("📥 Download BA + Audit", out, f"BA_AKA_{datetime.date.today()}.xlsx", use_container_width=True)
            except Exception as e: st.error(f"❌ {e}")
else:
    st.info("Upload RAB untuk mulai.", icon="⚖️"); st.markdown("<h1 style='text-align:center;font-size:80px'>⚖️</h1>", unsafe_allow_html=True)

st.caption("AKA v3.9.9 Production | Perpres 12/2021 & SE PUPR 07/2023")